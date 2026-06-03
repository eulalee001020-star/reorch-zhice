from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
ARTIFACTS = ROOT / "portfolio_artifacts"
BUILD = DIST / "portfolio_submission_20260603"
PACKAGE = BUILD / "ReOrch_Zhice_AI_Portfolio_20260603"
DOCX_PATH = DIST / "ReOrch_智策_AI产品作品集_20260603.docx"
PDF_PATH = DIST / "ReOrch_智策_AI产品作品集_20260603.pdf"
ZIP_PATH = DIST / "ReOrch_智策_AI作品集材料包_20260603.zip"

CN_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
PDF_FONT = "PortfolioCJK"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = LATIN_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6, line_spacing: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_paragraph(doc: Document, text: str = "", *, size: float = 11, bold: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(p, before=16, after=8, line_spacing=1.1)
        size, color = 16, BLUE
    elif level == 2:
        set_paragraph_spacing(p, before=12, after=6, line_spacing=1.1)
        size, color = 13, BLUE
    else:
        set_paragraph_spacing(p, before=8, after=4, line_spacing=1.1)
        size, color = 12, DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4, line_spacing=1.15)
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=10.5)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line_spacing=1.1)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold, color=color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_width(cell, widths[idx])
        shade_cell(cell, LIGHT_FILL)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.font.size = Pt(11)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("ReOrch 智策 AI 产品作品集")
    set_run_font(run, size=9, color=MUTED)


def build_docx() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=0, after=4, line_spacing=1.05)
    run = title.add_run("ReOrch 智策 AI 产品作品集")
    set_run_font(run, size=24, bold=True, color=RGBColor(0, 0, 0))

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=14, line_spacing=1.1)
    run = subtitle.add_run("工业异常调度决策 Copilot | AI 产品作品集 | 2026-06-03")
    set_run_font(run, size=11, color=MUTED)

    add_heading(doc, "1. 项目定位")
    add_paragraph(
        doc,
        "ReOrch 智策面向复杂离散制造中的异常调度决策：当设备故障、插单、物料延期、质量返工等事件打断原计划后，系统帮助计划员完成影响分析、策略选择、Top-K 候选方案比较、质量门校验、推荐解释、人工确认、受控回写和案例沉淀。",
    )
    add_paragraph(
        doc,
        "项目的核心判断不是让大模型直接自动排产，而是把 AI 放在可控工作流中。LLM/Agent 负责异常理解、规则候选、推荐解释、案例沉淀和偏好学习；求解器、硬约束、数字孪生验证、质量门、人工确认和审计链路负责生产责任。",
    )

    add_heading(doc, "2. 业务问题")
    add_paragraph(
        doc,
        "复杂制造企业通常已经有 ERP、MES、APS 和现场 Excel/人工沟通流程，但异常发生后的重排决策仍高度依赖计划员经验。设备故障、急单插入、物料延迟、质量返工和瓶颈资源冲突会同时影响交期、扰动范围、换线成本和执行风险。",
    )
    add_paragraph(
        doc,
        "ReOrch 的产品切口不是替换主系统，而是在主系统之上补“异常响应层”和“经验资产层”：把异常发生后的判断链路结构化、可解释化、可审计化，并让人工确认后的决策沉淀为可复用案例。",
    )

    add_heading(doc, "3. AI 能力证据")
    add_table(
        doc,
        ["能力", "项目证据"],
        [
            ["AI 产品定义", "把 AI 限定在语义、解释、规则候选和经验沉淀；不替代 APS/MES/LIMS 主系统"],
            ["Agent/Workflow", "Incident Intake、Constraint Compiler、Strategy Advisor、Explanation、Case Memory 等受控 Agent"],
            ["系统设计", "Incident -> Impact -> Solver -> Quality Gate -> Recommendation -> Confirmation -> Writeback -> Case Memory"],
            ["评测与质量门", "schema、source refs、硬约束、风险阈值、数字孪生 replay、失败样本库和人工确认"],
            ["工程落地", "FastAPI、React、OR-Tools、Docker Compose、mock integration、测试和 CI 验证材料"],
            ["商业判断", "先实验室试用和只读/shadow 验证，再进入客户现场；不宣称已生产上线"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "4. 工作流与演示")
    for item in [
        "Incident Intake -> Snapshot Lock -> Impact Analysis -> Strategy Advice -> Candidate Plan Generation -> Quality Gate -> Multi-objective Evaluation -> Recommendation Explanation -> Human Confirmation -> Controlled Writeback -> Case Memory。",
        "本地互动 demo：cp .env.example .env，然后运行 docker compose up --build，访问 http://localhost:3000，账号 planner / planner123。",
        "核心路径：登录 -> 决策工作台 -> 加载演示场景 -> 影响分析 -> Top-K 候选方案 -> 推荐解释 -> 人工确认 -> 受控回写 -> 案例库。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. 验证证据")
    add_table(
        doc,
        ["证据", "当前状态"],
        [
            ["后端测试", "pytest -q 已覆盖核心模型、API、Agent workflow、质量门和 demo sandbox"],
            ["前端构建", "React + TypeScript + Vite production build 已通过"],
            ["Demo 数据", "69 条 sandbox 记录，0 blocking error"],
            ["数字孪生验证", "已给出 source refs、replay/shadow 代理、风险分、阈值和审计包结构"],
            ["失败样本库", "明确不推荐、不自动写回、退回人工判断的条件"],
            ["LLM Agent 离线评测", "默认确定性降级可复现；真实 LLM Agent 支持模型、token、latency 和降级原因记录"],
        ],
        [2600, 6760],
    )

    add_heading(doc, "6. 关键边界")
    for item in [
        "当前项目证明的是 MVP、受控试用和 demo 级闭环，不等于客户生产系统正式上线。",
        "默认 demo 可在无外部 LLM API Key 的情况下复现；真实 LLM Agent 是可配置路径。",
        "系统不支持无人值守自动调度；生产回写必须经过人工确认、权限校验、回写预览和审计。",
        "数字孪生、synthetic package 和实验室 replay 是验证代理，不能替代客户现场数据、财务口径和上线验收。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. 材料索引")
    add_paragraph(
        doc,
        "关键入口为 README.md、service_ai_role_fit.md、prd_decision_workbench.md、mvp_delivery_plan.md、portfolio_proof_matrix.md、evaluation_guardrail_cases.md 和 failure_iteration_log.md。完整索引见 docs/portfolio/README.md。",
    )

    doc.save(DOCX_PATH)


def register_pdf_font() -> None:
    font_path = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")
    if font_path.exists():
        pdfmetrics.registerFont(TTFont(PDF_FONT, str(font_path)))
        return
    fallback = Path("/System/Library/Fonts/STHeiti Medium.ttc")
    if fallback.exists():
        pdfmetrics.registerFont(TTFont(PDF_FONT, str(fallback)))
        return
    raise FileNotFoundError("No CJK-capable font found for PDF generation.")


def pdf_styles() -> dict[str, ParagraphStyle]:
    return {
        "title": ParagraphStyle(
            "PortfolioTitle",
            fontName=PDF_FONT,
            fontSize=20,
            leading=25,
            spaceAfter=8,
            alignment=TA_LEFT,
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "PortfolioSubtitle",
            fontName=PDF_FONT,
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#595959"),
            spaceAfter=14,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "PortfolioH1",
            fontName=PDF_FONT,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "PortfolioBody",
            fontName=PDF_FONT,
            fontSize=9.8,
            leading=14,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "PortfolioBullet",
            fontName=PDF_FONT,
            fontSize=9.5,
            leading=13,
            leftIndent=12,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "table": ParagraphStyle(
            "PortfolioTable",
            fontName=PDF_FONT,
            fontSize=8.6,
            leading=11,
            wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "PortfolioTableHeader",
            fontName=PDF_FONT,
            fontSize=8.6,
            leading=11,
            textColor=colors.HexColor("#1F4D78"),
            wordWrap="CJK",
        ),
    }


def add_pdf_table(story, styles, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    data = [[Paragraph(header, styles["table_header"]) for header in headers]]
    for row in rows:
        data.append([Paragraph(value, styles["table"]) for value in row])
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D7DE")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))


def add_pdf_bullets(story, styles, items: list[str]) -> None:
    story.append(
        ListFlowable(
            [ListItem(Paragraph(item, styles["bullet"])) for item in items],
            bulletType="bullet",
            leftIndent=14,
            bulletFontName=PDF_FONT,
            bulletFontSize=8,
        )
    )
    story.append(Spacer(1, 4))


def build_pdf() -> None:
    register_pdf_font()
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title="ReOrch 智策 AI 产品作品集",
    )
    story = [
        Paragraph("ReOrch 智策 AI 产品作品集", styles["title"]),
        Paragraph("工业异常调度决策 Copilot | AI 产品作品集 | 2026-06-03", styles["subtitle"]),
    ]

    story.append(Paragraph("1. 项目定位", styles["h1"]))
    for text in [
        "ReOrch 智策面向复杂离散制造中的异常调度决策：当设备故障、插单、物料延期、质量返工等事件打断原计划后，系统帮助计划员完成影响分析、策略选择、Top-K 候选方案比较、质量门校验、推荐解释、人工确认、受控回写和案例沉淀。",
        "项目的核心判断不是让大模型直接自动排产，而是把 AI 放在可控工作流中。LLM/Agent 负责异常理解、规则候选、推荐解释、案例沉淀和偏好学习；求解器、硬约束、数字孪生验证、质量门、人工确认和审计链路负责生产责任。",
    ]:
        story.append(Paragraph(text, styles["body"]))

    story.append(Paragraph("2. 业务问题", styles["h1"]))
    for text in [
        "复杂制造企业通常已经有 ERP、MES、APS 和现场 Excel/人工沟通流程，但异常发生后的重排决策仍高度依赖计划员经验。设备故障、急单插入、物料延迟、质量返工和瓶颈资源冲突会同时影响交期、扰动范围、换线成本和执行风险。",
        "ReOrch 的产品切口不是替换主系统，而是在主系统之上补“异常响应层”和“经验资产层”：把异常发生后的判断链路结构化、可解释化、可审计化，并让人工确认后的决策沉淀为可复用案例。",
    ]:
        story.append(Paragraph(text, styles["body"]))

    story.append(Paragraph("3. AI 能力证据", styles["h1"]))
    add_pdf_table(
        story,
        styles,
        ["能力", "项目证据"],
        [
            ["AI 产品定义", "把 AI 限定在语义、解释、规则候选和经验沉淀；不替代 APS/MES/LIMS 主系统"],
            ["Agent/Workflow", "Incident Intake、Constraint Compiler、Strategy Advisor、Explanation、Case Memory 等受控 Agent"],
            ["系统设计", "Incident -> Impact -> Solver -> Quality Gate -> Recommendation -> Confirmation -> Writeback -> Case Memory"],
            ["评测与质量门", "schema、source refs、硬约束、风险阈值、数字孪生 replay、失败样本库和人工确认"],
            ["工程落地", "FastAPI、React、OR-Tools、Docker Compose、mock integration、测试和 CI 验证材料"],
            ["商业判断", "先实验室试用和只读/shadow 验证，再进入客户现场；不宣称已生产上线"],
        ],
        [120, 348],
    )

    story.append(Paragraph("4. 工作流与演示", styles["h1"]))
    add_pdf_bullets(
        story,
        styles,
        [
            "Incident Intake -> Snapshot Lock -> Impact Analysis -> Strategy Advice -> Candidate Plan Generation -> Quality Gate -> Multi-objective Evaluation -> Recommendation Explanation -> Human Confirmation -> Controlled Writeback -> Case Memory。",
            "本地互动 demo：cp .env.example .env，然后运行 docker compose up --build，访问 http://localhost:3000，账号 planner / planner123。",
            "核心路径：登录 -> 决策工作台 -> 加载演示场景 -> 影响分析 -> Top-K 候选方案 -> 推荐解释 -> 人工确认 -> 受控回写 -> 案例库。",
        ],
    )

    story.append(Paragraph("5. 验证证据", styles["h1"]))
    add_pdf_table(
        story,
        styles,
        ["证据", "当前状态"],
        [
            ["后端测试", "pytest -q 已覆盖核心模型、API、Agent workflow、质量门和 demo sandbox"],
            ["前端构建", "React + TypeScript + Vite production build 已通过"],
            ["Demo 数据", "69 条 sandbox 记录，0 blocking error"],
            ["数字孪生验证", "已给出 source refs、replay/shadow 代理、风险分、阈值和审计包结构"],
            ["失败样本库", "明确不推荐、不自动写回、退回人工判断的条件"],
            ["LLM Agent 离线评测", "默认确定性降级可复现；真实 LLM Agent 支持模型、token、latency 和降级原因记录"],
        ],
        [120, 348],
    )

    story.append(Paragraph("6. 关键边界", styles["h1"]))
    add_pdf_bullets(
        story,
        styles,
        [
            "当前证明的是 MVP、受控试用和 demo 级闭环，不等于客户生产系统正式上线。",
            "默认 demo 可在无外部 LLM API Key 的情况下复现；真实 LLM Agent 是可配置路径。",
            "系统不支持无人值守自动调度；生产回写必须经过人工确认、权限校验、回写预览和审计。",
            "数字孪生、synthetic package 和实验室 replay 是验证代理，不能替代客户现场数据和上线验收。",
        ],
    )

    story.append(Paragraph("7. 材料索引", styles["h1"]))
    story.append(
        Paragraph(
            "关键入口为 README.md、service_ai_role_fit.md、prd_decision_workbench.md、mvp_delivery_plan.md、portfolio_proof_matrix.md、evaluation_guardrail_cases.md 和 failure_iteration_log.md。完整索引见 docs/portfolio/README.md。",
            styles["body"],
        )
    )

    doc.build(story)


def should_skip(path: Path) -> bool:
    parts = set(path.parts)
    ignored_names = {
        "__pycache__",
        ".pytest_cache",
        ".ruff_cache",
        ".mypy_cache",
        ".hypothesis",
        ".git",
        "node_modules",
        "dist",
        "build",
        "htmlcov",
        "runtime",
    }
    if parts & ignored_names:
        return True
    return path.suffix in {".pyc", ".pyo", ".log", ".sqlite", ".sqlite3", ".db"}


def copy_file(src: Path, dst: Path) -> None:
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def copy_tree(src: Path, dst: Path) -> None:
    for path in src.rglob("*"):
        if should_skip(path):
            continue
        rel = path.relative_to(src)
        target = dst / rel
        if path.is_dir():
            target.mkdir(parents=True, exist_ok=True)
        else:
            copy_file(path, target)


def build_portfolio_readme() -> str:
    return """# ReOrch 智策 AI 产品作品集

ReOrch 智策是一个面向复杂离散制造异常重排的受控 AI 决策系统。项目展示了从业务问题识别、AI 能力边界、Agent workflow、求解器与质量门、人工确认、受控回写到案例沉淀的完整链路。

## 项目入口

- `README.md`：项目定位、作品集入口、demo 启动方式。
- `docs/portfolio/portfolio_brief.md`：AI 产品作品集摘要。
- `docs/portfolio/service_ai_role_fit.md`：服务领域 AI 产品能力映射。
- `docs/portfolio/service_ai_transfer_note.md`：服务领域 AI 迁移说明。
- `docs/product/prd_decision_workbench.md`：标准 PRD 示例。
- `docs/project/mvp_delivery_plan.md`：MVP 交付计划与项目推进。
- `docs/portfolio/portfolio_proof_matrix.md`：作品集证明材料矩阵。
- `docs/portfolio/ai_native_pm_capability_map.md`：AI Native 产品经理能力映射。
- `docs/portfolio/industrial_ai_copilot_solution.md`：工业 AI Copilot 方案说明。
- `docs/portfolio/business_process_flow.md`：业务流程图、泳道图和状态机。
- `docs/portfolio/prototype_logic.md`：工作台信息架构、页面逻辑和状态处理。
- `docs/portfolio/metric_system.md`：North Star、分层指标和采集路径。
- `docs/portfolio/evaluation_guardrail_cases.md`：Guardrail 用例与验证标准。
- `docs/portfolio/failure_iteration_log.md`：失败案例与迭代记录。
- `docs/portfolio/cost_latency_deployment_boundary.md`：成本、延迟与部署边界。
- `docs/portfolio/personal_contribution.md`：个人贡献说明。
- `docs/portfolio/project_report_materials.md`：项目汇报结构、答辩问题和材料包索引。
- `docs/portfolio/product_portfolio.md`：项目背景、AI 取舍、产品设计和证据索引。
- `docs/portfolio/workflow_prompts_io.md`：Agent workflow、prompt 模板、输入输出样例。
- `docs/portfolio/trust_quality_gate.md` 与 `docs/validation/failure_case_library.md`：可信性、失败样本和上线边界。

## 本地互动 demo

```bash
cp .env.example .env
docker compose up --build
```

访问 `http://localhost:3000`，账号 `planner / planner123`。

核心路径：登录 -> 决策工作台 -> 加载演示场景 -> 影响分析 -> Top-K 候选方案 -> 推荐解释 -> 人工确认 -> 受控回写 -> 案例库。

## 关键边界

- 本项目证明 MVP、受控试用和 demo 级闭环，不宣称客户生产系统已经正式上线。
- LLM/Agent 负责语义、解释、规则候选和经验沉淀；求解器、质量门、人工确认和审计负责生产责任。
- 默认 demo 不依赖外部 LLM API Key；真实 LLM Agent 是可配置路径。
"""


def build_package() -> None:
    if BUILD.exists():
        shutil.rmtree(BUILD)
    PACKAGE.mkdir(parents=True, exist_ok=True)

    root_files = [
        "README.md",
        ".env.example",
        "docker-compose.yml",
        "Dockerfile",
        "Makefile",
        "pyproject.toml",
        "alembic.ini",
    ]
    for name in root_files:
        src = ROOT / name
        if src.exists():
            copy_file(src, PACKAGE / name)

    if PDF_PATH.exists():
        copy_file(PDF_PATH, PACKAGE / PDF_PATH.name)

    for directory in [
        "app",
        "alembic",
        "docs",
        "demo",
        "benchmark",
        ".github/workflows",
        "frontend/src",
    ]:
        src = ROOT / directory
        if src.exists():
            copy_tree(src, PACKAGE / directory)

    for name in [
        "frontend/package.json",
        "frontend/package-lock.json",
        "frontend/index.html",
        "frontend/vite.config.ts",
        "frontend/tsconfig.json",
        "frontend/nginx.conf",
        "frontend/Dockerfile",
    ]:
        src = ROOT / name
        if src.exists():
            copy_file(src, PACKAGE / name)

    (PACKAGE / "PORTFOLIO_README.md").write_text(build_portfolio_readme(), encoding="utf-8")


def build_zip() -> None:
    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in PACKAGE.rglob("*"):
            if path.is_file():
                archive.write(path, path.relative_to(BUILD))


def main() -> None:
    DIST.mkdir(exist_ok=True)
    build_docx()
    build_pdf()
    build_package()
    build_zip()
    ARTIFACTS.mkdir(exist_ok=True)
    for artifact in [DOCX_PATH, PDF_PATH, ZIP_PATH]:
        copy_file(artifact, ARTIFACTS / artifact.name)
    print(DOCX_PATH)
    print(PDF_PATH)
    print(ZIP_PATH)


if __name__ == "__main__":
    main()
